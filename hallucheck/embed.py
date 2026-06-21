"""Render a verification packet into linked documents.

Two stdlib renderers (Markdown, HTML) and two optional ones (DOCX via
``python-docx``, PDF via ``reportlab``; install the ``[docs]`` extra). Every
renderer emits, per authority:

  * the **citation** with a stable anchor / bookmark,
  * **source links** an attorney can follow (official source, Google Scholar,
    CourtListener, web-archive snapshot) and clearly-labeled subscription / bar
    portals,
  * the **source text** itself (proof the authority exists and reads as cited) with
    its SHA-256,
  * a **treatment** block (the attorney's Shepardize result, if recorded), and
  * **cross-links** to related authorities — in DOCX/PDF these are real internal
    links/bookmarks, so a brief's citation jumps to the section showing the cited
    text, and a noted negative treatment links onward to the next authority.

Every document is stamped with the library + corpus disclaimer.
"""
from __future__ import annotations

import html

from .disclaimer import SHORT_DISCLAIMER


def _status(entry: dict) -> str:
    if entry["dead_link"]:
        return "DEAD LINK — source URL unreachable; verify another way"
    if entry["resolved"]:
        return "text retrieved (verify it is current & on point)"
    return "UNRESOLVED — not in the trusted index; verify manually before relying"


def render(packet: dict, fmt: str, path: str | None = None):
    """Render to ``fmt`` (md|markdown|html|docx|pdf). String formats return the
    string (and also write it when ``path`` is given); docx/pdf require ``path``."""
    fmt = fmt.lower()
    if fmt in ("md", "markdown"):
        out = to_markdown(packet)
    elif fmt == "html":
        out = to_html(packet)
    elif fmt == "docx":
        if not path:
            raise ValueError("docx output requires a path")
        return to_docx(packet, path)
    elif fmt == "pdf":
        if not path:
            raise ValueError("pdf output requires a path")
        return to_pdf(packet, path)
    else:
        raise ValueError(f"unknown format {fmt!r}")
    if path:
        import pathlib
        pathlib.Path(path).write_text(out, encoding="utf-8")
    return out


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #
def _md_links(entry: dict) -> list[str]:
    rows = []
    for ln in entry["sources"]["links"]:
        if ln.get("url"):
            rows.append(f"  - [{ln['label']}]({ln['url']})")
        elif ln.get("view_url"):                       # web-archive snapshot
            rows.append(f"  - {ln['label']}: [view]({ln['view_url']}) · "
                        f"[save snapshot]({ln['save_url']})")
    for p in entry["sources"]["portals"]:
        rows.append(f"  - {p['label']} — [portal]({p['portal_url']}) "
                    f"_(login required; search: `{p['query']}`)_")
    return rows


def to_markdown(packet: dict) -> str:
    L: list[str] = [f"# {packet['title']}", ""]
    L.append("> **" + SHORT_DISCLAIMER + "**")
    for line in packet["disclaimer"].splitlines():
        L.append("> " + line if line.strip() else ">")
    L += ["",
          f"*Adapter:* `{packet.get('adapter')}` · *scope:* `{packet.get('scope')}` · "
          f"*generated:* {packet['generated_at']} · *config:* `{packet.get('config_digest')}`",
          ""]
    c = packet["counts"]
    L.append(f"**{c['total']} authorities** — resolved {c['resolved']}, "
             f"unresolved {c['unresolved']}, dead links {c['dead_links']}, "
             f"treatment recorded {c['treated']}.")
    if packet["unverified"]:
        L.append("")
        L.append("> ⚠️ **Verify manually (did not resolve):** "
                 + ", ".join(f"`{u}`" for u in packet["unverified"]))
    L += ["", "## Table of authorities", ""]
    for i, e in enumerate(packet["entries"], 1):
        L.append(f"{i}. [{e['cite']}](#{e['anchor']}) — {e['title'] or ''} "
                 f"_({e['kind']})_ — {_status(e)}")
    L.append("")
    for e in packet["entries"]:
        L.append(f'<a id="{e["anchor"]}"></a>')
        L.append(f"## {e['cite']}")
        if e["title"]:
            L.append(f"*{e['title']}* — kind: {e['kind']}")
        L.append("")
        L.append("**Read & verify:**")
        L += _md_links(e)
        L.append("")
        if e["text"]:
            L.append(f"**Source text** (SHA-256 `{e['text_sha256']}`):")
            L.append("")
            for line in e["text"].splitlines() or [""]:
                L.append("> " + line)
        else:
            L.append(f"> **{_status(e)}**")
        L.append("")
        cl = e.get("courtlistener") or {}
        if cl.get("found"):
            L.append(f"**CourtListener:** {cl.get('case_name') or ''} — "
                     f"{cl.get('court') or ''}, {cl.get('date') or ''}")
            if cl.get("snippet"):
                L.append("> " + " ".join(cl["snippet"].split()))
            L.append("")
        t = e["treatment"]
        L.append(f"**Treatment:** `{t['status']}`"
                 + (f" — {t['note']}" if t.get("note") else "")
                 + (f" — reviewed by {t['reviewed_by']}" if t.get("reviewed_by") else ""))
        for a in t.get("authorities", []):
            L.append(f"  - see {_md_treat_auth(a)}")
        if e["related"]:
            L.append("**Related authorities:** "
                     + ", ".join(f"[{r['cite']}](#{r['anchor']})" for r in e["related"]))
        L.append("")
    return "\n".join(L).rstrip() + "\n"


def _md_treat_auth(a: dict) -> str:
    label = a.get("label") or a.get("cite") or a.get("url") or "authority"
    if a.get("anchor"):
        return f"[{label}](#{a['anchor']})"
    if a.get("url"):
        return f"[{label}]({a['url']})"
    return label


# --------------------------------------------------------------------------- #
# HTML
# --------------------------------------------------------------------------- #
_CSS = """
body{font:16px/1.5 -apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;
max-width:54rem;margin:2rem auto;padding:0 1rem;color:#111}
.disclaimer{background:#fff8e1;border:1px solid #e0c04a;border-radius:6px;
padding:.75rem 1rem;margin:1rem 0;font-size:.95rem;white-space:pre-wrap}
.disclaimer b{color:#9a3b00}
.meta{color:#555;font-size:.85rem}
section{border-top:1px solid #ddd;padding-top:1rem;margin-top:1.5rem}
.cite{font-weight:700}
blockquote{border-left:3px solid #bbb;margin:.5rem 0;padding:.25rem .9rem;
background:#fafafa;white-space:pre-wrap}
.sha{font-family:ui-monospace,Menlo,Consolas,monospace;font-size:.78rem;color:#555}
.status-unresolved{color:#9a0007;font-weight:700}
.status-dead{color:#9a0007;font-weight:700}
.treat{background:#f1f6ff;border:1px solid #cfe0ff;border-radius:6px;padding:.5rem .8rem}
ul.links{margin:.3rem 0}
.login{color:#777;font-size:.85rem}
""".strip()


def _h(s) -> str:
    return html.escape(str(s if s is not None else ""))


def _html_links(entry: dict) -> str:
    items = []
    for ln in entry["sources"]["links"]:
        if ln.get("url"):
            items.append(f'<li><a href="{_h(ln["url"])}" target="_blank" '
                         f'rel="noopener">{_h(ln["label"])}</a></li>')
        elif ln.get("view_url"):
            items.append(f'<li>{_h(ln["label"])}: '
                         f'<a href="{_h(ln["view_url"])}" target="_blank" rel="noopener">view</a> · '
                         f'<a href="{_h(ln["save_url"])}" target="_blank" rel="noopener">save snapshot</a></li>')
    for p in entry["sources"]["portals"]:
        items.append(f'<li><a href="{_h(p["portal_url"])}" target="_blank" '
                     f'rel="noopener">{_h(p["label"])}</a> '
                     f'<span class="login">(login required; search: '
                     f'<code>{_h(p["query"])}</code>)</span></li>')
    return '<ul class="links">' + "".join(items) + "</ul>"


def to_html(packet: dict) -> str:
    p = packet
    out = ["<!doctype html><html lang=en><head><meta charset=utf-8>",
           f"<meta name=viewport content='width=device-width,initial-scale=1'>",
           f"<title>{_h(p['title'])}</title><style>{_CSS}</style></head><body>",
           f"<h1>{_h(p['title'])}</h1>",
           f'<div class="disclaimer"><b>{_h(SHORT_DISCLAIMER)}</b>\n\n{_h(p["disclaimer"])}</div>',
           f'<p class="meta">Adapter <code>{_h(p.get("adapter"))}</code> · scope '
           f'<code>{_h(p.get("scope"))}</code> · generated {_h(p["generated_at"])} · '
           f'config <code>{_h(p.get("config_digest"))}</code></p>']
    c = p["counts"]
    out.append(f"<p><b>{c['total']} authorities</b> — resolved {c['resolved']}, "
               f"unresolved {c['unresolved']}, dead links {c['dead_links']}, "
               f"treatment recorded {c['treated']}.</p>")
    if p["unverified"]:
        out.append('<p class="status-unresolved">⚠️ Verify manually (did not resolve): '
                   + ", ".join(f"<code>{_h(u)}</code>" for u in p["unverified"]) + "</p>")
    out.append("<h2>Table of authorities</h2><ol>")
    for e in p["entries"]:
        out.append(f'<li><a href="#{_h(e["anchor"])}">{_h(e["cite"])}</a> — '
                   f'{_h(e["title"])} <em>({_h(e["kind"])})</em> — {_h(_status(e))}</li>')
    out.append("</ol>")
    for e in p["entries"]:
        out.append(f'<section id="{_h(e["anchor"])}">')
        out.append(f'<h2 class="cite">{_h(e["cite"])}</h2>')
        if e["title"]:
            out.append(f'<p class="meta">{_h(e["title"])} — kind: {_h(e["kind"])}</p>')
        out.append("<p><b>Read &amp; verify:</b></p>" + _html_links(e))
        if e["text"]:
            out.append(f'<details open><summary>Source text '
                       f'<span class="sha">SHA-256 {_h(e["text_sha256"])}</span></summary>'
                       f'<blockquote>{_h(e["text"])}</blockquote></details>')
        else:
            cls = "status-dead" if e["dead_link"] else "status-unresolved"
            out.append(f'<p class="{cls}">{_h(_status(e))}</p>')
        cl = e.get("courtlistener") or {}
        if cl.get("found"):
            out.append(f'<p class="meta"><b>CourtListener:</b> {_h(cl.get("case_name"))} — '
                       f'{_h(cl.get("court"))}, {_h(cl.get("date"))}</p>')
            if cl.get("snippet"):
                out.append(f'<blockquote>{_h(" ".join(cl["snippet"].split()))}</blockquote>')
        t = e["treatment"]
        treat = (f'<div class="treat"><b>Treatment:</b> <code>{_h(t["status"])}</code>'
                 + (f' — {_h(t["note"])}' if t.get("note") else "")
                 + (f' — reviewed by {_h(t["reviewed_by"])}' if t.get("reviewed_by") else ""))
        auths = t.get("authorities", [])
        if auths:
            treat += "<ul>" + "".join(f"<li>see {_html_treat_auth(a)}</li>" for a in auths) + "</ul>"
        out.append(treat + "</div>")
        if e["related"]:
            out.append("<p><b>Related authorities:</b> " + ", ".join(
                f'<a href="#{_h(r["anchor"])}">{_h(r["cite"])}</a>' for r in e["related"]) + "</p>")
        out.append("</section>")
    out.append("</body></html>")
    return "".join(out)


def _html_treat_auth(a: dict) -> str:
    label = a.get("label") or a.get("cite") or a.get("url") or "authority"
    if a.get("anchor"):
        return f'<a href="#{_h(a["anchor"])}">{_h(label)}</a>'
    if a.get("url"):
        return f'<a href="{_h(a["url"])}" target="_blank" rel="noopener">{_h(label)}</a>'
    return _h(label)


# --------------------------------------------------------------------------- #
# DOCX (optional: python-docx)
# --------------------------------------------------------------------------- #
def to_docx(packet: dict, path: str):
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.shared import RGBColor, Pt
    except Exception as e:                                  # pragma: no cover
        raise RuntimeError("DOCX output needs the [docs] extra: pip install "
                           "'llm-hallucination-checker[docs]'") from e

    state = {"bm": 0}

    def external_link(par, url, text):
        r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
        h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
        run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        h.append(run); par._p.append(h)

    def internal_link(par, anchor, text):
        h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), anchor)
        run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        h.append(run); par._p.append(h)

    def bookmark(par, name):
        state["bm"] += 1
        s = OxmlElement("w:bookmarkStart")
        s.set(qn("w:id"), str(state["bm"])); s.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(state["bm"]))
        par._p.insert(0, s); par._p.append(end)

    doc = Document()
    doc.add_heading(packet["title"], level=0)
    d = doc.add_paragraph()
    run = d.add_run(SHORT_DISCLAIMER + "\n\n" + packet["disclaimer"])
    run.italic = True; run.font.color.rgb = RGBColor(0x9A, 0x3B, 0x00); run.font.size = Pt(9)
    meta = doc.add_paragraph()
    meta.add_run(f"Adapter: {packet.get('adapter')}  ·  scope: {packet.get('scope')}  ·  "
                 f"generated: {packet['generated_at']}  ·  config: {packet.get('config_digest')}"
                 ).font.size = Pt(8)

    doc.add_heading("Table of authorities", level=1)
    for i, e in enumerate(packet["entries"], 1):
        par = doc.add_paragraph(style="List Number")
        internal_link(par, e["anchor"], e["cite"])
        par.add_run(f" — {e['title'] or ''} ({e['kind']}) — {_status(e)}")

    for e in packet["entries"]:
        head = doc.add_heading(e["cite"], level=1)
        bookmark(head, e["anchor"])
        if e["title"]:
            doc.add_paragraph(f"{e['title']} — kind: {e['kind']}").runs[0].italic = True
        doc.add_paragraph().add_run("Read & verify:").bold = True
        for ln in e["sources"]["links"]:
            p = doc.add_paragraph(style="List Bullet")
            if ln.get("url"):
                external_link(p, ln["url"], ln["label"])
            elif ln.get("view_url"):
                p.add_run(ln["label"] + ": ")
                external_link(p, ln["view_url"], "view"); p.add_run(" · ")
                external_link(p, ln["save_url"], "save snapshot")
        for pl in e["sources"]["portals"]:
            p = doc.add_paragraph(style="List Bullet")
            external_link(p, pl["portal_url"], pl["label"])
            p.add_run(f"  (login required; search: {pl['query']})").italic = True

        if e["text"]:
            doc.add_paragraph().add_run(f"Source text (SHA-256 {e['text_sha256']}):").bold = True
            q = doc.add_paragraph(e["text"]); q.style = "Quote"
        else:
            doc.add_paragraph().add_run(_status(e)).bold = True

        cl = e.get("courtlistener") or {}
        if cl.get("found"):
            clp = doc.add_paragraph(); clp.add_run("CourtListener: ").bold = True
            clp.add_run(f"{cl.get('case_name') or ''} — {cl.get('court') or ''}, "
                        f"{cl.get('date') or ''}")
            if cl.get("snippet"):
                doc.add_paragraph(" ".join(cl["snippet"].split())).style = "Quote"

        t = e["treatment"]
        tp = doc.add_paragraph()
        tp.add_run("Treatment: ").bold = True
        tp.add_run(t["status"] + (f" — {t['note']}" if t.get("note") else "")
                   + (f" — reviewed by {t['reviewed_by']}" if t.get("reviewed_by") else ""))
        for a in t.get("authorities", []):
            ap = doc.add_paragraph(style="List Bullet"); ap.add_run("see ")
            _docx_treat_auth(ap, a, internal_link, external_link)
        if e["related"]:
            rp = doc.add_paragraph(); rp.add_run("Related authorities: ").bold = True
            for j, r in enumerate(e["related"]):
                if j:
                    rp.add_run(", ")
                internal_link(rp, r["anchor"], r["cite"])

    doc.save(path)
    return path


def _docx_treat_auth(par, a, internal_link, external_link):
    label = a.get("label") or a.get("cite") or a.get("url") or "authority"
    if a.get("anchor"):
        internal_link(par, a["anchor"], label)
    elif a.get("url"):
        external_link(par, a["url"], label)
    else:
        par.add_run(label)


# --------------------------------------------------------------------------- #
# PDF (optional: reportlab) — outline bookmarks + internal & external links
# --------------------------------------------------------------------------- #
def to_pdf(packet: dict, path: str):
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Flowable)
    except Exception as e:                                  # pragma: no cover
        raise RuntimeError("PDF output needs the [docs] extra: pip install "
                           "'llm-hallucination-checker[docs]'") from e

    import xml.sax.saxutils as su

    def esc(s):
        return su.escape(str(s if s is not None else ""))

    class _Outline(Flowable):
        """Zero-height flowable: registers a PDF outline (sidebar) entry + a named
        page destination for the section that follows it."""
        def __init__(self, key, title, level=0):
            super().__init__()
            self.key, self.title, self.level = key, title, level
            self.width = self.height = 0

        def draw(self):
            self.canv.bookmarkPage(self.key)
            self.canv.addOutlineEntry(self.title[:120], self.key, self.level, False)

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=13)
    small = ParagraphStyle("small", parent=body, fontSize=8, textColor=HexColor("#555555"))
    quote = ParagraphStyle("quote", parent=body, leftIndent=14, backColor=HexColor("#f5f5f5"),
                           borderColor=HexColor("#bbbbbb"), borderWidth=0.5,
                           borderPadding=4, spaceBefore=4, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=14, spaceBefore=14)
    disc = ParagraphStyle("disc", parent=small, textColor=HexColor("#9a3b00"),
                          backColor=HexColor("#fff8e1"), borderColor=HexColor("#e0c04a"),
                          borderWidth=0.5, borderPadding=6, spaceAfter=10)

    story = [Paragraph(esc(packet["title"]), styles["Title"]),
             Paragraph("<b>" + esc(SHORT_DISCLAIMER) + "</b><br/><br/>"
                       + esc(packet["disclaimer"]).replace("\n", "<br/>"), disc),
             Paragraph(f"Adapter {esc(packet.get('adapter'))} · scope "
                       f"{esc(packet.get('scope'))} · generated {esc(packet['generated_at'])} "
                       f"· config {esc(packet.get('config_digest'))}", small),
             Spacer(1, 8), Paragraph("Table of authorities", h1)]
    for i, e in enumerate(packet["entries"], 1):
        story.append(Paragraph(
            f'{i}. <a href="#{esc(e["anchor"])}" color="blue">{esc(e["cite"])}</a> — '
            f'{esc(e["title"])} ({esc(e["kind"])}) — {esc(_status(e))}', body))
    for e in packet["entries"]:
        story.append(_Outline(e["anchor"] + "-top", e["cite"], 0))
        story.append(Paragraph(f'<a name="{esc(e["anchor"])}"/>{esc(e["cite"])}', h1))
        if e["title"]:
            story.append(Paragraph(f'<i>{esc(e["title"])}</i> — kind: {esc(e["kind"])}', small))
        story.append(Paragraph("<b>Read &amp; verify:</b>", body))
        for ln in e["sources"]["links"]:
            if ln.get("url"):
                story.append(Paragraph(
                    f'• <a href="{esc(ln["url"])}" color="blue">{esc(ln["label"])}</a>', body))
            elif ln.get("view_url"):
                story.append(Paragraph(
                    f'• {esc(ln["label"])}: <a href="{esc(ln["view_url"])}" color="blue">view</a> '
                    f'· <a href="{esc(ln["save_url"])}" color="blue">save snapshot</a>', body))
        for pl in e["sources"]["portals"]:
            story.append(Paragraph(
                f'• <a href="{esc(pl["portal_url"])}" color="blue">{esc(pl["label"])}</a> '
                f'<font size=8 color="#777777">(login required; search: {esc(pl["query"])})</font>',
                body))
        if e["text"]:
            story.append(Paragraph(f'<b>Source text</b> <font size=8 color="#555555">SHA-256 '
                                   f'{esc(e["text_sha256"])}</font>', body))
            story.append(Paragraph(esc(e["text"]).replace("\n", "<br/>"), quote))
        else:
            story.append(Paragraph(f'<b><font color="#9a0007">{esc(_status(e))}</font></b>', body))
        cl = e.get("courtlistener") or {}
        if cl.get("found"):
            story.append(Paragraph(f"<b>CourtListener:</b> {esc(cl.get('case_name'))} — "
                                   f"{esc(cl.get('court'))}, {esc(cl.get('date'))}", body))
            if cl.get("snippet"):
                story.append(Paragraph(esc(" ".join(cl["snippet"].split())), quote))
        t = e["treatment"]
        tline = (f'<b>Treatment:</b> {esc(t["status"])}'
                 + (f' — {esc(t["note"])}' if t.get("note") else "")
                 + (f' — reviewed by {esc(t["reviewed_by"])}' if t.get("reviewed_by") else ""))
        story.append(Paragraph(tline, body))
        for a in t.get("authorities", []):
            story.append(Paragraph("• see " + _pdf_treat_auth(a, esc), body))
        if e["related"]:
            rel = ", ".join(f'<a href="#{esc(r["anchor"])}" color="blue">{esc(r["cite"])}</a>'
                            for r in e["related"])
            story.append(Paragraph(f"<b>Related authorities:</b> {rel}", body))
        story.append(Spacer(1, 6))

    doc = BaseDocTemplate(path, pagesize=LETTER, title=packet["title"],
                          leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                          topMargin=0.9 * inch, bottomMargin=0.9 * inch)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="f")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame])])
    doc.build(story)
    return path


def _pdf_treat_auth(a: dict, esc) -> str:
    label = a.get("label") or a.get("cite") or a.get("url") or "authority"
    if a.get("anchor"):
        return f'<a href="#{esc(a["anchor"])}" color="blue">{esc(label)}</a>'
    if a.get("url"):
        return f'<a href="{esc(a["url"])}" color="blue">{esc(label)}</a>'
    return esc(label)
